import docx
from docx import Document
import aspose.words as aw
import os
import sys

def formatter(folder):
    
    # create a list of all files in directory to apply this to
    relative_paths = []
    for root, dirs, files in os.walk(folder):
        for file in files:
            relative_paths.append(os.path.relpath(os.path.join(root, file), folder))
            
    # loop through every file           
    for x in relative_paths:           
                
        filepath = (x)
        print(filepath)

        # convert docx into txt file here
        doc = docx.Document(filepath)
        new_file = "example.txt"


        with open(new_file, "w") as file:
            # Loop through all paragraphs in the document and write the text to the file
            for paragraph in doc.paragraphs:
                file.write(paragraph.text + '\n')



        # formats txt file as requested
        with open('example.txt', 'r') as f:
            text = f.read()
            paragraphs = text.split('\n\n')
            new_text = ''
            for p in paragraphs:
                lines = p.split('\n')
                if len(lines) > 2:
                    new_text += '\n'.join(lines[2:]) + '\n\n'
            with open('new_example.txt', 'w') as f2:
                f2.write(new_text)  
                    
        # delete tempfile 
        os.remove("example.txt")
        


        # convert back from txt to docx
        with open('new_example.txt', 'r') as file:
            text = file.read()

        document = Document()
        paragraph = document.add_paragraph(text)

        newname =  "formatted" + filepath
        document.save(newname)
        
        os.remove("new_example.txt")

droppedFile = sys.argv[1]           
formatter(droppedFile)
#C:\Users\maisa\Desktop\testfolder